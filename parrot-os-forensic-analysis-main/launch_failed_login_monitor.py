"""
File: launch_failed_login_monitor.py
Authors:
    - Harshitha S M
    - Deepti Bhat
    - Nandish H R
"""
#!/usr/bin/python3

import subprocess
import sys
import logging
import os
from logging.handlers import RotatingFileHandler
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# Define the log file and rotation parameters
log_file_path = '/var/log/Forensic_Report/loginreport.log'
max_bytes = 1024 * 1024  # 1 MB
backup_count = 5

os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

# Create a rotating file handler
handler = RotatingFileHandler(log_file_path, maxBytes=max_bytes, backupCount=backup_count)
formatter = logging.Formatter('%(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

text_to_append = """Linux Forensic Analysis Report.

Executive Summary
Based on the set thresholds, we have observed failed login attempts that exceed the allowed count. Please find the details below.

Incident Details
"""

def clear_text_file(file_path):
    try:
        with open(file_path, 'w') as file:
            pass
        print(f"Text file '{file_path}' has been cleared.")
    except IOError as e:
        print(f"Error clearing text file '{file_path}': {e}")

def write_output_to_file(output):
    try:
        with open('loginreport.txt', 'a') as file:
            file.write(output)
        print("Output written to report.txt")
        if "End of Report" in output:
            try:
                generate_styled_report('loginreport.txt', 'loginreport_styled.docx')
                subprocess.run(f"python3 send_mail.py login {user}", shell=True)
            except FileNotFoundError:
                print("Error: The script file was not found.")
            except Exception as e:
                print("An error occurred:", e)
            clear_text_file("loginreport.txt")
    except IOError as e:
        print(f"Error writing to report.txt: {e}")


def generate_styled_report(input_path, output_path):
    with open(input_path, 'r') as file:
        lines = file.readlines()

    doc = Document()

    # Title
    title = doc.add_heading('Linux Forensic Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)  # Make the title bigger

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = doc.add_paragraph('Based on the set thresholds, we have observed failed login attempts that exceed the allowed count. Please find the details below.')
    exec_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incident Details
    doc.add_heading('Incident Details', level=1)

    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith('User associated with failed logins'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('Start timestamp'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('End timestamp'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('Failed login attempts'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('Failed login from IP address'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('Location'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped_line.startswith('------'):
            p = doc.add_paragraph(stripped_line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT   

    # End of Report
    doc.add_heading('End of Report', level=1)
    end_report = doc.add_paragraph('This report summarizes the analysis of failed login attempts within the specified time frame. Please review the details provided above for security assessments and necessary actions.')
    end_report.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.text = "Lhedge-Linux Services"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save(output_path)
    print(f"Styled report saved to '{output_path}'")


def execute_bash_script(script_path):
    try:
        result = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True, text=True)
        while True: 
            output = result.stdout.readline()
            if output == '' and result.poll() is not None:
                break
            if output:
                logger.info(output.strip())
                print(output.strip())
                write_output_to_file(output)
    except subprocess.CalledProcessError as e:
        print("Error executing bash script:", e)

script_path = 'detect_failed_login.sh'
lff = sys.argv[1]
lfd = sys.argv[2]
user = sys.argv[3]
command = f"bash {script_path} {lff} {lfd}"
execute_bash_script(script_path)
