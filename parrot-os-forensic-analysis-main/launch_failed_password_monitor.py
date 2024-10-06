"""
File: launch_failed_password_monitor.py
Authors:
    - Harshitha S M
    - Deepti Bhat
    - Nandish H R
"""
#!/usr/bin/python3

import subprocess
import sys
import logging
import os
from logging.handlers import RotatingFileHandler
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# Define the log file and rotation parameters
log_file_path = '/var/log/Forensic_Report/passChange_report.log'
max_bytes = 1024 * 1024  # 1 MB
backup_count = 5

os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

# Create a rotating file handler
handler = RotatingFileHandler(log_file_path, maxBytes=max_bytes, backupCount=backup_count)
formatter = logging.Formatter('%(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

text_to_append = """Linux Forensic Analysis Report.

Executive Summary
Based on the set thresholds, we have observed password changes that exceed the allowed count. Please find the details below.

Incident Details
"""

def clear_text_file(file_path):
    try:
        with open(file_path, 'w') as file:
            pass
        print(f"Text file '{file_path}' has been cleared.")
    except IOError as e:
        print(f"Error clearing text file '{file_path}': {e}")

def append_text_to_file(text, file_path):
    try:
        with open(file_path, 'a') as file:
            file.write(text)
            file.write("\n")
        print(f"Text appended to '{file_path}'")
    except IOError as e:
        print(f"Error appending text to '{file_path}': {e}")

def write_output_to_file(output):
    try:
        with open('passreport.txt', 'a') as file:
            file.write(output)
        print("Output written to passreport.txt")
        if "running" in output:
            generate_styled_report('passreport.txt', 'passreport_styled.docx')
            try:
                subprocess.run(f"python3 send_mail.py pass {user}", shell=True)
            except FileNotFoundError:
                print("Error: The script file was not found.")
            except Exception as e:
                print("An error occurred:", e)
            clear_text_file("passreport.txt")
            append_text_to_file(text_to_append, "passreport.txt")
    except IOError as e:
        print(f"Error writing to passreport.txt: {e}")

def generate_styled_report(input_path, output_path):
    with open(input_path, 'r') as file:
        lines = file.readlines()

    doc = Document()

    # Title
    title = doc.add_heading('Linux Forensic Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)  # Make the title bigger

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = doc.add_paragraph('Based on the set thresholds, we have observed password changes that exceed the allowed count. Please find the details below.')
    exec_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incident Details
    doc.add_heading('Incident Details', level=1)

    # Skip the initial text in the file
    start_line = 0
    for i, line in enumerate(lines):
        if "Incident Details" in line:
            start_line = i + 1
            break

    for line in lines[start_line:]:
        # Skip the line if it contains 'running'
        if 'running' not in line:
            if "Password changes detected" in line:
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.strip():
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # End of Report
    doc.add_heading('End of Report', level=1)
    end_report = doc.add_paragraph('This report summarizes the analysis of password changes within the specified time frame. Please review the details provided above for security assessments and necessary actions.')
    end_report.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.text = "Lhedge-Linux Services"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save(output_path)
    print(f"Styled report saved to '{output_path}'")

def execute_bash_script(script_path):
    try:
        result = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True, text=True)
        while True: 
            output = result.stdout.readline()
            if output == '' and result.poll() is not None:
                break
            if output:
                print(output.strip())
                logger.info(output.strip())
                write_output_to_file(output)
    except subprocess.CalledProcessError as e:
        print("Error executing bash script:", e)

script_path = 'detect_failed_password.sh'
pcf = sys.argv[1]
pcd = sys.argv[2]
user = sys.argv[3]
command = f"bash {script_path} {pcf} {pcd}"
execute_bash_script(script_path)
